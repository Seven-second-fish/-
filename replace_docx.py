import os
import sys
from docx import Document

def replace_text_in_paragraphs(paragraphs, old_string, new_string):
    for paragraph in paragraphs:
        if old_string in paragraph.text:
            paragraph.text = paragraph.text.replace(old_string, new_string)

def replace_text_in_tables(tables, old_string, new_string):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, old_string, new_string)

def replace_in_document(file_path, old_string, new_string):
    try:
        doc = Document(file_path)
        replace_text_in_paragraphs(doc.paragraphs, old_string, new_string)
        replace_text_in_tables(doc.tables, old_string, new_string)
        doc.save(file_path)
        print(f"Processed: {file_path}")
    except Exception as e:
        print(f"Error processing file {file_path}: {e}")

def replace_in_docx_files(root_dir, old_string, new_string):
    for dirpath, dirnames, filenames in os.walk(root_dir):
        for filename in filenames:
            if filename.endswith('.docx') and not filename.startswith('.') and not filename.startswith('~'):
                file_path = os.path.join(dirpath, filename)
                replace_in_document(file_path, old_string, new_string)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python replace_docx.py <old_string> <new_string>")
        sys.exit(1)

    folder_path = os.getcwd()  # 获取当前脚本所在文件夹的路径
    old_string = sys.argv[1]
    new_string = sys.argv[2]

    replace_in_docx_files(folder_path, old_string, new_string)

